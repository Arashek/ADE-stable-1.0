from typing import Dict, List, Any, Optional, Union
from datetime import datetime
import logging
from dataclasses import dataclass
from enum import Enum
import hashlib
import base64
import io
import os
from PIL import Image
import numpy as np
import torch
from transformers import (
    AutoProcessor, AutoModel, AutoTokenizer, AutoModelForSeq2SeqLM,
    AutoModelForImageClassification, AutoModelForCausalLM
)
import soundfile as sf
import librosa
import json
import cv2
import pytesseract
from pdf2image import convert_from_bytes
import docx
import ast
import re
from typing import Dict, List, Any, Optional, Union, Tuple
import networkx as nx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import DBSCAN
import torchaudio
import torchvision.transforms as transforms
from transformers import pipeline

from .memory_manager import MemoryEntry, MemoryType

logger = logging.getLogger(__name__)

class InputType(Enum):
    """Types of input modalities"""
    TEXT = "text"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    DOCUMENT = "document"
    CODE = "code"
    STRUCTURED_DATA = "structured_data"

@dataclass
class ProcessedInput:
    """Represents processed input data"""
    content: str
    embedding: List[float]
    metadata: Dict[str, Any]
    content_hash: str
    summary: Optional[str] = None
    compressed_content: Optional[str] = None
    modality: str = "text"
    original_type: InputType = InputType.TEXT
    original_data: Optional[bytes] = None
    processed_data: Optional[Dict[str, Any]] = None

class MultiModalProcessor:
    """Processes different types of input modalities"""
    
    def __init__(
        self,
        model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
        device: str = "cuda" if torch.cuda.is_available() else "cpu"
    ):
        """Initialize multimodal processor
        
        Args:
            model_name: Name of the model to use for embeddings
            device: Device to run the model on
        """
        self.device = device
        self.model_name = model_name
        
        # Initialize models and processors
        self._initialize_models()
        
        # Supported modalities and their processors
        self.modality_processors = {
            InputType.TEXT: self._process_text,
            InputType.IMAGE: self._process_image,
            InputType.AUDIO: self._process_audio,
            InputType.VIDEO: self._process_video,
            InputType.DOCUMENT: self._process_document,
            InputType.CODE: self._process_code,
            InputType.STRUCTURED_DATA: self._process_structured_data
        }
        
    def _initialize_models(self):
        """Initialize required models and processors"""
        try:
            # Initialize text embedding model
            self.text_processor = AutoProcessor.from_pretrained(self.model_name)
            self.text_model = AutoModel.from_pretrained(self.model_name).to(self.device)
            
            # Initialize image processing model
            self.image_processor = AutoProcessor.from_pretrained("google/vit-base-patch16-224")
            self.image_model = AutoModel.from_pretrained("google/vit-base-patch16-224").to(self.device)
            
            # Initialize audio processing model
            self.audio_processor = AutoProcessor.from_pretrained("facebook/wav2vec2-base")
            self.audio_model = AutoModel.from_pretrained("facebook/wav2vec2-base").to(self.device)
            
            # Initialize text summarization model
            self.summarizer = pipeline(
                "summarization",
                model="facebook/bart-large-cnn",
                device=0 if torch.cuda.is_available() else -1
            )
            
            # Initialize image description model
            self.image_describer = pipeline(
                "image-to-text",
                model="nlpconnect/vit-gpt2-image-captioning",
                device=0 if torch.cuda.is_available() else -1
            )
            
            # Initialize speech recognition model
            self.speech_recognizer = pipeline(
                "automatic-speech-recognition",
                model="facebook/wav2vec2-base-960h",
                device=0 if torch.cuda.is_available() else -1
            )
            
            # Initialize code analysis model
            self.code_analyzer = AutoModelForCausalLM.from_pretrained(
                "microsoft/codebert-base"
            ).to(self.device)
            self.code_tokenizer = AutoTokenizer.from_pretrained("microsoft/codebert-base")
            
            # Initialize document processing models
            self.document_processor = AutoModelForSeq2SeqLM.from_pretrained(
                "facebook/bart-large-cnn"
            ).to(self.device)
            self.document_tokenizer = AutoTokenizer.from_pretrained("facebook/bart-large-cnn")
            
        except Exception as e:
            logger.error(f"Failed to initialize models: {str(e)}")
            raise
            
    async def process_input(
        self,
        input_data: Union[str, bytes, Dict[str, Any]],
        input_type: InputType,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedInput:
        """Process input data of any supported modality
        
        Args:
            input_data: Input data to process
            input_type: Type of input data
            metadata: Additional metadata
            
        Returns:
            Processed input data
        """
        try:
            # Get appropriate processor
            processor = self.modality_processors.get(input_type)
            if not processor:
                raise ValueError(f"Unsupported input type: {input_type}")
                
            # Process input
            processed_data = await processor(input_data)
            
            # Generate content hash
            content_hash = self._generate_content_hash(processed_data)
            
            # Create processed input
            processed_input = ProcessedInput(
                content=processed_data["content"],
                embedding=processed_data["embedding"],
                metadata=metadata or {},
                content_hash=content_hash,
                summary=processed_data.get("summary"),
                compressed_content=processed_data.get("compressed_content"),
                modality=input_type.value,
                original_type=input_type,
                original_data=input_data if isinstance(input_data, bytes) else None,
                processed_data=processed_data
            )
            
            return processed_input
            
        except Exception as e:
            logger.error(f"Failed to process input: {str(e)}")
            raise
            
    async def _process_text(
        self,
        text: str
    ) -> Dict[str, Any]:
        """Process text input
        
        Args:
            text: Text to process
            
        Returns:
            Processed text data
        """
        try:
            # Generate embedding
            inputs = self.text_processor(
                text,
                return_tensors="pt",
                padding=True,
                truncation=True,
                max_length=512
            ).to(self.device)
            
            with torch.no_grad():
                outputs = self.text_model(**inputs)
                embedding = outputs.last_hidden_state.mean(dim=1).cpu().numpy()[0]
                
            # Generate summary if text is long
            summary = None
            if len(text.split()) > 100:
                summary = self._generate_summary(text)
                
            return {
                "content": text,
                "embedding": embedding.tolist(),
                "summary": summary
            }
            
        except Exception as e:
            logger.error(f"Failed to process text: {str(e)}")
            raise
            
    async def _process_image(
        self,
        image_data: bytes
    ) -> Dict[str, Any]:
        """Process image input
        
        Args:
            image_data: Image data to process
            
        Returns:
            Processed image data
        """
        try:
            # Convert bytes to PIL Image
            image = Image.open(io.BytesIO(image_data))
            
            # Generate embedding
            inputs = self.image_processor(
                image,
                return_tensors="pt"
            ).to(self.device)
            
            with torch.no_grad():
                outputs = self.image_model(**inputs)
                embedding = outputs.last_hidden_state.mean(dim=1).cpu().numpy()[0]
                
            # Generate text description
            description = self._generate_image_description(image)
            
            # Compress image
            compressed_data = self._compress_image(image)
            
            return {
                "content": description,
                "embedding": embedding.tolist(),
                "summary": description,
                "compressed_content": compressed_data,
                "image_metadata": {
                    "size": image.size,
                    "mode": image.mode,
                    "format": image.format
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to process image: {str(e)}")
            raise
            
    async def _process_audio(
        self,
        audio_data: bytes
    ) -> Dict[str, Any]:
        """Process audio input
        
        Args:
            audio_data: Audio data to process
            
        Returns:
            Processed audio data
        """
        try:
            # Convert bytes to numpy array
            audio_array, sample_rate = sf.read(io.BytesIO(audio_data))
            
            # Generate embedding
            inputs = self.audio_processor(
                audio_array,
                sampling_rate=sample_rate,
                return_tensors="pt"
            ).to(self.device)
            
            with torch.no_grad():
                outputs = self.audio_model(**inputs)
                embedding = outputs.last_hidden_state.mean(dim=1).cpu().numpy()[0]
                
            # Generate transcription
            transcription = self._transcribe_audio(audio_array, sample_rate)
            
            # Extract features
            features = self._extract_audio_features(audio_array, sample_rate)
            
            return {
                "content": transcription,
                "embedding": embedding.tolist(),
                "summary": transcription,
                "audio_metadata": {
                    "duration": len(audio_array) / sample_rate,
                    "sample_rate": sample_rate,
                    "features": features
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to process audio: {str(e)}")
            raise
            
    async def _process_video(
        self,
        video_data: bytes
    ) -> Dict[str, Any]:
        """Process video input
        
        Args:
            video_data: Video data to process
            
        Returns:
            Processed video data
        """
        try:
            # Extract frames and audio
            frames, audio = self._extract_video_components(video_data)
            
            # Process frames
            frame_embeddings = []
            frame_descriptions = []
            for frame in frames:
                frame_data = await self._process_image(frame)
                frame_embeddings.append(frame_data["embedding"])
                frame_descriptions.append(frame_data["content"])
                
            # Process audio
            audio_data = await self._process_audio(audio)
            
            # Combine embeddings
            combined_embedding = np.mean([frame_embeddings, audio_data["embedding"]], axis=0)
            
            # Generate video description
            description = self._generate_video_description(frame_descriptions, audio_data["content"])
            
            return {
                "content": description,
                "embedding": combined_embedding.tolist(),
                "summary": description,
                "video_metadata": {
                    "num_frames": len(frames),
                    "frame_descriptions": frame_descriptions,
                    "audio_transcription": audio_data["content"]
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to process video: {str(e)}")
            raise
            
    async def _process_document(
        self,
        document_data: bytes
    ) -> Dict[str, Any]:
        """Process document input
        
        Args:
            document_data: Document data to process
            
        Returns:
            Processed document data
        """
        try:
            # Extract content and metadata
            text, metadata = self._extract_document_content(document_data)
            
            # Process text
            text_data = await self._process_text(text)
            
            return {
                "content": text_data["content"],
                "embedding": text_data["embedding"],
                "summary": text_data["summary"],
                "document_metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to process document: {str(e)}")
            raise
            
    async def _process_code(
        self,
        code_data: str
    ) -> Dict[str, Any]:
        """Process code input
        
        Args:
            code_data: Code to process
            
        Returns:
            Processed code data
        """
        try:
            # Process code
            processed_code = self._process_code_content(code_data)
            
            # Generate embedding
            code_embedding = await self._process_text(processed_code["content"])
            
            return {
                "content": processed_code["content"],
                "embedding": code_embedding["embedding"],
                "summary": processed_code["summary"],
                "code_metadata": processed_code["metadata"]
            }
            
        except Exception as e:
            logger.error(f"Failed to process code: {str(e)}")
            raise
            
    async def _process_structured_data(
        self,
        data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process structured data input
        
        Args:
            data: Structured data to process
            
        Returns:
            Processed structured data
        """
        try:
            # Convert to text representation
            text_representation = self._convert_structured_to_text(data)
            
            # Process text
            text_data = await self._process_text(text_representation)
            
            return {
                "content": text_data["content"],
                "embedding": text_data["embedding"],
                "summary": text_data["summary"],
                "structured_metadata": {
                    "schema": self._infer_schema(data),
                    "type": type(data).__name__
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to process structured data: {str(e)}")
            raise
            
    def _generate_content_hash(self, data: Dict[str, Any]) -> str:
        """Generate content hash for data
        
        Args:
            data: Data to hash
            
        Returns:
            Content hash
        """
        content = json.dumps(data, sort_keys=True)
        return hashlib.sha256(content.encode()).hexdigest()
        
    def _generate_summary(self, text: str) -> str:
        """Generate summary for text
        
        Args:
            text: Text to summarize
            
        Returns:
            Generated summary
        """
        try:
            # Split text into chunks if too long
            max_chunk_length = 1024
            chunks = [text[i:i + max_chunk_length] for i in range(0, len(text), max_chunk_length)]
            
            summaries = []
            for chunk in chunks:
                # Generate summary for each chunk
                summary = self.summarizer(
                    chunk,
                    max_length=130,
                    min_length=30,
                    do_sample=False
                )[0]["summary_text"]
                summaries.append(summary)
            
            # Combine summaries
            return " ".join(summaries)
            
        except Exception as e:
            logger.error(f"Failed to generate summary: {str(e)}")
            return text[:200] + "..."  # Fallback to simple truncation
        
    def _generate_image_description(self, image: Image.Image) -> str:
        """Generate description for image
        
        Args:
            image: Image to describe
            
        Returns:
            Generated description
        """
        try:
            # Convert PIL Image to RGB if needed
            if image.mode != "RGB":
                image = image.convert("RGB")
            
            # Generate description using image-to-text model
            description = self.image_describer(
                image,
                max_length=50,
                num_beams=5,
                early_stopping=True
            )[0]["generated_text"]
            
            return description
            
        except Exception as e:
            logger.error(f"Failed to generate image description: {str(e)}")
            return "Image description unavailable"
        
    def _compress_image(self, image: Image.Image) -> str:
        """Compress image
        
        Args:
            image: Image to compress
            
        Returns:
            Compressed image data as base64 string
        """
        try:
            # Convert PIL Image to RGB if needed
            if image.mode != "RGB":
                image = image.convert("RGB")
            
            # Create a BytesIO object to store the compressed image
            img_byte_arr = io.BytesIO()
            
            # Save image with compression
            image.save(
                img_byte_arr,
                format="JPEG",
                quality=85,  # Adjust quality (0-100)
                optimize=True
            )
            
            # Convert to base64 string
            img_byte_arr.seek(0)
            compressed_data = base64.b64encode(img_byte_arr.getvalue()).decode()
            
            return compressed_data
            
        except Exception as e:
            logger.error(f"Failed to compress image: {str(e)}")
            return ""
        
    def _transcribe_audio(
        self,
        audio_array: np.ndarray,
        sample_rate: int
    ) -> str:
        """Transcribe audio
        
        Args:
            audio_array: Audio data
            sample_rate: Sample rate
            
        Returns:
            Transcription
        """
        try:
            # Convert numpy array to torch tensor
            audio_tensor = torch.from_numpy(audio_array).float()
            
            # Resample if needed (model expects 16kHz)
            if sample_rate != 16000:
                resampler = torchaudio.transforms.Resample(sample_rate, 16000)
                audio_tensor = resampler(audio_tensor)
            
            # Transcribe using speech recognition model
            transcription = self.speech_recognizer(
                audio_tensor,
                sampling_rate=16000,
                max_length=448,
                chunk_length_s=30
            )["text"]
            
            return transcription
            
        except Exception as e:
            logger.error(f"Failed to transcribe audio: {str(e)}")
            return "Audio transcription unavailable"
        
    def _extract_audio_features(
        self,
        audio_array: np.ndarray,
        sample_rate: int
    ) -> Dict[str, Any]:
        """Extract features from audio
        
        Args:
            audio_array: Audio data
            sample_rate: Sample rate
            
        Returns:
            Extracted features
        """
        try:
            features = {}
            
            # Extract basic features
            features["duration"] = len(audio_array) / sample_rate
            features["sample_rate"] = sample_rate
            
            # Extract MFCC features
            mfccs = librosa.feature.mfcc(y=audio_array, sr=sample_rate, n_mfcc=13)
            features["mfcc_mean"] = np.mean(mfccs, axis=1).tolist()
            features["mfcc_std"] = np.std(mfccs, axis=1).tolist()
            
            # Extract spectral features
            spec_cent = librosa.feature.spectral_centroid(y=audio_array, sr=sample_rate)
            spec_bw = librosa.feature.spectral_bandwidth(y=audio_array, sr=sample_rate)
            spec_flat = librosa.feature.spectral_flatness(y=audio_array)
            spec_rolloff = librosa.feature.spectral_rolloff(y=audio_array, sr=sample_rate)
            
            features["spectral_centroid"] = np.mean(spec_cent).tolist()
            features["spectral_bandwidth"] = np.mean(spec_bw).tolist()
            features["spectral_flatness"] = np.mean(spec_flat).tolist()
            features["spectral_rolloff"] = np.mean(spec_rolloff).tolist()
            
            # Extract rhythm features
            tempo, _ = librosa.beat.beat_track(y=audio_array, sr=sample_rate)
            features["tempo"] = tempo
            
            # Extract zero-crossing rate
            zcr = librosa.feature.zero_crossing_rate(audio_array)
            features["zero_crossing_rate"] = np.mean(zcr).tolist()
            
            return features
            
        except Exception as e:
            logger.error(f"Failed to extract audio features: {str(e)}")
            return {}
        
    def _extract_video_components(
        self,
        video_data: bytes
    ) -> tuple:
        """Extract frames and audio from video
        
        Args:
            video_data: Video data
            
        Returns:
            Tuple of (frames, audio)
        """
        try:
            # Create temporary file for video data
            with open("temp_video.mp4", "wb") as f:
                f.write(video_data)
            
            # Open video file
            cap = cv2.VideoCapture("temp_video.mp4")
            
            # Extract frames
            frames = []
            frame_interval = int(cap.get(cv2.CAP_PROP_FPS))  # Extract one frame per second
            frame_count = 0
            
            while cap.isOpened():
                ret, frame = cap.read()
                if not ret:
                    break
                    
                if frame_count % frame_interval == 0:
                    # Convert BGR to RGB
                    frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    # Convert to PIL Image
                    frame_pil = Image.fromarray(frame_rgb)
                    frames.append(frame_pil)
                    
                frame_count += 1
                
            cap.release()
            
            # Extract audio
            audio_data = None
            try:
                # Use moviepy to extract audio
                from moviepy.editor import VideoFileClip
                video = VideoFileClip("temp_video.mp4")
                audio = video.audio
                if audio is not None:
                    # Save audio to temporary file
                    audio.write_audiofile("temp_audio.wav")
                    # Read audio data
                    with open("temp_audio.wav", "rb") as f:
                        audio_data = f.read()
                video.close()
            except Exception as e:
                logger.warning(f"Failed to extract audio from video: {str(e)}")
            
            # Clean up temporary files
            os.remove("temp_video.mp4")
            if os.path.exists("temp_audio.wav"):
                os.remove("temp_audio.wav")
            
            return frames, audio_data
            
        except Exception as e:
            logger.error(f"Failed to extract video components: {str(e)}")
            return [], None
        
    def _generate_video_description(
        self,
        frame_descriptions: List[str],
        audio_transcription: str
    ) -> str:
        """Generate description for video
        
        Args:
            frame_descriptions: List of frame descriptions
            audio_transcription: Audio transcription
            
        Returns:
            Generated description
        """
        try:
            # Combine frame descriptions
            frame_summary = " ".join(frame_descriptions)
            
            # Generate video description using document processor
            inputs = self.document_tokenizer(
                f"Frames: {frame_summary}\nAudio: {audio_transcription}",
                return_tensors="pt",
                max_length=1024,
                truncation=True
            ).to(self.device)
            
            outputs = self.document_processor.generate(
                **inputs,
                max_length=150,
                num_beams=5,
                early_stopping=True
            )
            
            description = self.document_tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            return description
            
        except Exception as e:
            logger.error(f"Failed to generate video description: {str(e)}")
            return "Video description unavailable"
        
    def _extract_document_content(
        self,
        document_data: bytes
    ) -> tuple:
        """Extract content and metadata from document
        
        Args:
            document_data: Document data
            
        Returns:
            Tuple of (content, metadata)
        """
        try:
            # Create temporary file for document data
            with open("temp_doc", "wb") as f:
                f.write(document_data)
            
            content = ""
            metadata = {}
            
            # Try to detect document type and extract content
            try:
                # Try PDF
                images = convert_from_bytes(document_data)
                for image in images:
                    content += pytesseract.image_to_string(image) + "\n"
                metadata["type"] = "pdf"
                metadata["pages"] = len(images)
                
            except Exception as e:
                try:
                    # Try DOCX
                    doc = docx.Document("temp_doc")
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    metadata["type"] = "docx"
                    metadata["paragraphs"] = len(doc.paragraphs)
                    
                except Exception as e:
                    # Try plain text
                    with open("temp_doc", "r", encoding="utf-8") as f:
                        content = f.read()
                    metadata["type"] = "text"
            
            # Clean up temporary file
            os.remove("temp_doc")
            
            # Extract additional metadata
            metadata["length"] = len(content)
            metadata["lines"] = content.count("\n")
            metadata["words"] = len(content.split())
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Failed to extract document content: {str(e)}")
            return "", {}
        
    def _process_code_content(self, code: str) -> Dict[str, Any]:
        """Process code content
        
        Args:
            code: Code to process
            
        Returns:
            Processed code data
        """
        try:
            # Try to parse code as AST
            try:
                tree = ast.parse(code)
                # Extract function and class names
                functions = [node.name for node in ast.walk(tree) if isinstance(node, ast.FunctionDef)]
                classes = [node.name for node in ast.walk(tree) if isinstance(node, ast.ClassDef)]
                
                # Extract imports
                imports = []
                for node in ast.walk(tree):
                    if isinstance(node, ast.Import):
                        imports.extend(alias.name for alias in node.names)
                    elif isinstance(node, ast.ImportFrom):
                        imports.extend(f"{node.module}.{alias.name}" for alias in node.names)
                
            except SyntaxError:
                # If not valid Python code, treat as plain text
                functions = []
                classes = []
                imports = []
            
            # Generate code embedding
            inputs = self.code_tokenizer(
                code,
                return_tensors="pt",
                max_length=512,
                truncation=True,
                padding=True
            ).to(self.device)
            
            with torch.no_grad():
                outputs = self.code_analyzer(**inputs)
                embedding = outputs.last_hidden_state.mean(dim=1).cpu().numpy()[0]
            
            # Generate code summary
            summary = self._generate_summary(code)
            
            # Extract code metrics
            lines = code.split("\n")
            metrics = {
                "lines": len(lines),
                "non_empty_lines": len([line for line in lines if line.strip()]),
                "functions": len(functions),
                "classes": len(classes),
                "imports": len(imports),
                "complexity": self._calculate_code_complexity(code)
            }
            
            return {
                "content": code,
                "embedding": embedding.tolist(),
                "summary": summary,
                "metadata": {
                    "functions": functions,
                    "classes": classes,
                    "imports": imports,
                    "metrics": metrics
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to process code: {str(e)}")
            return {
                "content": code,
                "embedding": [],
                "summary": "Code processing failed",
                "metadata": {}
            }
            
    def _calculate_code_complexity(self, code: str) -> int:
        """Calculate code complexity
        
        Args:
            code: Code to analyze
            
        Returns:
            Complexity score
        """
        try:
            # Count control structures
            control_structures = [
                r"\bif\b",
                r"\bfor\b",
                r"\bwhile\b",
                r"\bexcept\b",
                r"\bwith\b",
                r"\breturn\b",
                r"\bbreak\b",
                r"\bcontinue\b"
            ]
            
            complexity = 0
            for structure in control_structures:
                complexity += len(re.findall(structure, code))
                
            return complexity
            
        except Exception as e:
            logger.error(f"Failed to calculate code complexity: {str(e)}")
            return 0
        
    def _convert_structured_to_text(self, data: Dict[str, Any]) -> str:
        """Convert structured data to text
        
        Args:
            data: Structured data
            
        Returns:
            Text representation
        """
        return json.dumps(data, indent=2)
        
    def _infer_schema(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Infer schema from structured data
        
        Args:
            data: Structured data
            
        Returns:
            Inferred schema
        """
        try:
            schema = {}
            
            def infer_type(value: Any) -> str:
                """Infer type of a value"""
                if isinstance(value, bool):
                    return "boolean"
                elif isinstance(value, int):
                    return "integer"
                elif isinstance(value, float):
                    return "float"
                elif isinstance(value, str):
                    return "string"
                elif isinstance(value, list):
                    if value:
                        return f"array[{infer_type(value[0])}]"
                    return "array"
                elif isinstance(value, dict):
                    return "object"
                else:
                    return "unknown"
            
            def process_value(value: Any, path: str = "") -> None:
                """Process a value and update schema"""
                if isinstance(value, dict):
                    for key, val in value.items():
                        new_path = f"{path}.{key}" if path else key
                        schema[new_path] = {
                            "type": infer_type(val),
                            "required": True
                        }
                        process_value(val, new_path)
                elif isinstance(value, list):
                    if value:
                        item_type = infer_type(value[0])
                        schema[path] = {
                            "type": f"array[{item_type}]",
                            "required": True
                        }
                        process_value(value[0], path)
            
            # Process the data
            for key, value in data.items():
                schema[key] = {
                    "type": infer_type(value),
                    "required": True
                }
                process_value(value, key)
            
            return schema
            
        except Exception as e:
            logger.error(f"Failed to infer schema: {str(e)}")
            return {} 